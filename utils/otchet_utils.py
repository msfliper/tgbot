import asyncio
import datetime


from database.crud.get import get_last_72_hours_reports
from enums import ReportAdminStatus
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from loader import bot
from aiogram.types import FSInputFile


def create_otchet_1() -> str:
    reports = get_last_72_hours_reports()
    total_number_reports = len(reports)
    reviewed_reports_number = len(list(filter(lambda report: report.admin_status == ReportAdminStatus.RESOLVED,
                                              reports)))
    not_reviewed_reports_number = len(list(filter(lambda report: report.admin_status == ReportAdminStatus.NOT_HIRING,
                                                  reports)))
    cannot_be_considered = len(list(filter(lambda report: report.admin_status == ReportAdminStatus.REJECT,
                                           reports)))
    datatime_now = datetime.datetime.now()
    datatime_3_days_ago = datatime_now - datetime.timedelta(days=3)

    otchet_1 = docx.Document()

    paragraph1 = otchet_1.add_paragraph()
    run1 = paragraph1.add_run('Отчет')
    run1.font.size = Pt(30)
    paragraph1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph2 = otchet_1.add_paragraph()
    run2 = paragraph2.add_run(f"За период с {datatime_3_days_ago.day}.{datatime_3_days_ago.month}.{datatime_3_days_ago.year} "
                              f"по {datatime_now.day}.{datatime_now.month}.{datatime_now.year}")
    run2.font.size = Pt(18)

    paragraph3 = otchet_1.add_paragraph()
    run3 = paragraph3.add_run(f"Всего жалоб было отправлено: {total_number_reports}\n"
                              f"Всего рассмотрено жалоб: {reviewed_reports_number}\n"
                              f"Всего не рассмотрено жалоб: {not_reviewed_reports_number}\n"
                              f"Всего жалоб, которые не могут быть рассмотрены: {cannot_be_considered}\n")

    paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    file_name = f"othchet1_{datatime_now.day}_{datatime_now.month}_{datatime_now.year}.docx"
    otchet_1.save(file_name)
    return file_name


def create_otchet_2() -> str:
    reports = get_last_72_hours_reports()
    reports_without_reject = [report for report in reports
                              if report.admin_status != ReportAdminStatus.REJECT]
    reports_without_solution = [report for report in reports_without_reject
                                if report.solution_at is None]
    reports_with_solution = [report for report in reports_without_reject
                             if report.solution_at]
    overdue_reports_number = (
            len([report for report in reports_with_solution
                 if (report.solution_at-report.created_at).days >= 1]) +
            len([report for report in reports_without_solution
                 if (datetime.datetime.now() - report.created_at).days >= 1]))
    resolved_reports_number = len([report for report in reports_with_solution
                                   if (report.solution_at-report.created_at).days < 1])
    datatime_now = datetime.datetime.now()
    datatime_3_days_ago = datatime_now - datetime.timedelta(days=3)

    otchet_2 = docx.Document()

    paragraph1 = otchet_2.add_paragraph()
    run1 = paragraph1.add_run('Отчет')
    run1.font.size = Pt(30)
    paragraph1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph2 = otchet_2.add_paragraph()
    run2 = paragraph2.add_run(
        f"За период с {datatime_3_days_ago.day}.{datatime_3_days_ago.month}.{datatime_3_days_ago.year} "
        f"по {datatime_now.day}.{datatime_now.month}.{datatime_now.year}")
    run2.font.size = Pt(18)
    paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph3 = otchet_2.add_paragraph()
    run3 = paragraph3.add_run(f"Всего жалоб, которые не успели проверить за сутки: {overdue_reports_number}\n"
                              f"Всего провереных вовремя жалоб: {resolved_reports_number}")
    file_name = f"othchet2_{datatime_now.day}_{datatime_now.month}_{datatime_now.year}.docx"
    otchet_2.save(file_name)
    return file_name


async def send_othcet(chat_id: int):
    othchet1_filename = create_otchet_1()
    othchet2_filename = create_otchet_2()
    await bot.send_message(text="Отчет за 3 дня",
                           chat_id=chat_id)
    await bot.send_document(chat_id=chat_id,
                            document=FSInputFile(path=f"./{othchet1_filename}"))
    await bot.send_document(chat_id=chat_id,
                            document=FSInputFile(path=f"./{othchet2_filename}"))

